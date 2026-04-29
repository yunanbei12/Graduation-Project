from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name, font_size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "黑体" if style_name != "Title" else "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style_name != "Title" else "宋体")
        style.font.size = Pt(font_size)
        style.font.bold = True


def apply_run_style(run, bold: bool = False, code: bool = False) -> None:
    if code:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10.5)
    else:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)
    run.bold = bold


def add_rich_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(24)

    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            apply_run_style(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            apply_run_style(run, code=True)
        else:
            run = paragraph.add_run(part)
            apply_run_style(run)


def flush_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        if set(stripped.replace("|", "").replace("-", "").replace(" ", "").replace(":", "")) == set():
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        cells = table.add_row().cells
        for c_index in range(col_count):
            text = row[c_index] if c_index < len(row) else ""
            cells[c_index].text = text
            for paragraph in cells[c_index].paragraphs:
                for run in paragraph.runs:
                    apply_run_style(run, bold=(r_index == 0))


def flush_code_block(doc: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        run = paragraph.add_run(line.rstrip("\n"))
        apply_run_style(run, code=True)


def add_image(doc: Document, markdown_dir: Path, alt_text: str, raw_path: str) -> None:
    image_path = Path(raw_path)
    if not image_path.is_absolute():
        image_path = (markdown_dir / image_path).resolve()
    if not image_path.exists():
        add_rich_paragraph(doc, f"[图片缺失] {raw_path}")
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))

    if alt_text.strip():
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(alt_text.strip())
        apply_run_style(run)


def render_markdown(input_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    in_code_block = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    title_written = False

    def flush_pending_table() -> None:
        nonlocal table_lines
        if table_lines:
            flush_table(doc, table_lines)
            table_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_pending_table()
            if in_code_block:
                flush_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            table_lines.append(line)
            continue
        else:
            flush_pending_table()

        if not line.strip():
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line.strip())
        if image_match:
            add_image(doc, input_path.parent, image_match.group(1), image_match.group(2))
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if not title_written:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                apply_run_style(run, bold=True)
                run.font.size = Pt(18)
                title_written = True
            else:
                p = doc.add_paragraph(style="Heading 1")
                run = p.add_run(text)
                apply_run_style(run, bold=True)
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(line[3:].strip())
            apply_run_style(run, bold=True)
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(line[4:].strip())
            apply_run_style(run, bold=True)
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            run = p.add_run(line[5:].strip())
            apply_run_style(run, bold=True)
            continue

        if re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            apply_run_style(run)
            continue

        add_rich_paragraph(doc, line)

    flush_pending_table()
    if code_lines:
        flush_code_block(doc, code_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: export_thesis_docx.py <input.md> <output.docx>")
        return 1
    input_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    render_markdown(input_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
