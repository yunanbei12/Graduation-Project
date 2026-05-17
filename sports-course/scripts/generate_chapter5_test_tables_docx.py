from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = "第五章_系统测试用例表.docx"


TEST_TABLES = [
    (
        "表5-1 用户预约下单功能测试用例",
        [
            ("测试用例编号", "TC01"),
            ("测试用例名称", "用户预约下单功能测试"),
            (
                "测试步骤",
                "1. 用户登录小程序端进入课程详情页。\n"
                "2. 选择课程规格、上课时间或场次信息。\n"
                "3. 选择可用优惠券并提交订单。\n"
                "4. 检查订单页面及订单列表结果。",
            ),
            (
                "预期目标",
                "1. 系统能够正确展示课程信息与价格信息。\n"
                "2. 提交订单时能够完成参数校验。\n"
                "3. 使用优惠券后订单金额计算正确。\n"
                "4. 下单成功后生成对应订单记录。\n"
                "5. 订单列表能够显示新生成的订单。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-2 团课场次校验功能测试用例",
        [
            ("测试用例编号", "TC02"),
            ("测试用例名称", "团课场次校验功能测试"),
            (
                "测试步骤",
                "1. 用户进入团课详情页。\n"
                "2. 选择一个可预约场次提交订单。\n"
                "3. 再次使用多个账号对同一场次重复报名。\n"
                "4. 检查满员情况下系统提示信息。",
            ),
            (
                "预期目标",
                "1. 系统能够正确校验团课场次是否存在。\n"
                "2. 系统能够正确校验场次剩余名额。\n"
                "3. 名额充足时允许正常下单。\n"
                "4. 名额不足或场次不可用时给出明确提示。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-3 课程排课功能测试用例",
        [
            ("测试用例编号", "TC03"),
            ("测试用例名称", "课程排课功能测试"),
            (
                "测试步骤",
                "1. 管理员进入后台排课管理页面。\n"
                "2. 选择课程、教练、上课时间和上课地点并提交排课。\n"
                "3. 查看排课列表与教练排课信息。\n"
                "4. 重复提交同时间段排课信息进行校验。",
            ),
            (
                "预期目标",
                "1. 系统能够正确保存排课信息。\n"
                "2. 排课列表能够正常展示课程、教练、时间和地点等内容。\n"
                "3. 教练排课信息能够同步更新。\n"
                "4. 对于明显冲突或无效数据能够进行必要提示。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-4 私教销课功能测试用例",
        [
            ("测试用例编号", "TC04"),
            ("测试用例名称", "私教销课功能测试"),
            (
                "测试步骤",
                "1. 管理员进入销课管理页面。\n"
                "2. 查询某用户的私教订单或课包信息。\n"
                "3. 执行一次销课操作。\n"
                "4. 查看剩余课时、销课记录和订单状态变化。",
            ),
            (
                "预期目标",
                "1. 系统能够正确识别可销课订单。\n"
                "2. 销课成功后剩余课时自动扣减。\n"
                "3. 系统能够生成对应销课记录。\n"
                "4. 订单相关状态与课时信息能够同步更新。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-5 团课签到与结课功能测试用例",
        [
            ("测试用例编号", "TC05"),
            ("测试用例名称", "团课签到与结课功能测试"),
            (
                "测试步骤",
                "1. 管理员进入团课签到页面。\n"
                "2. 选择已排课的团课场次。\n"
                "3. 对报名用户执行签到操作。\n"
                "4. 完成签到后检查订单状态和课次记录。",
            ),
            (
                "预期目标",
                "1. 系统能够正常展示场次报名名单。\n"
                "2. 签到操作能够成功提交。\n"
                "3. 签到后对应用户上课记录能够正确保存。\n"
                "4. 课程完成后订单状态能够按业务规则更新。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-6 AI咨询应答功能测试用例",
        [
            ("测试用例编号", "TC06"),
            ("测试用例名称", "AI咨询应答功能测试"),
            (
                "测试步骤",
                "1. 用户进入小程序 AI 客服页面。\n"
                "2. 输入课程咨询问题，如课程价格、课程类型或上课安排。\n"
                "3. 发送咨询内容并查看系统回复。\n"
                "4. 多次输入不同类型问题观察回复情况。",
            ),
            (
                "预期目标",
                "1. 系统能够正常接收用户咨询内容。\n"
                "2. 对常见问题能够返回较为准确的回复结果。\n"
                "3. 回复内容能够正常展示在会话列表中。\n"
                "4. 系统整体响应过程稳定，无明显异常。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-7 转人工处理功能测试用例",
        [
            ("测试用例编号", "TC07"),
            ("测试用例名称", "转人工处理功能测试"),
            (
                "测试步骤",
                "1. 用户在 AI 客服页面输入“转人工”等请求。\n"
                "2. 提交会话内容并查看返回结果。\n"
                "3. 管理员进入后台查看 AI 会话管理页面。\n"
                "4. 检查对应会话是否生成转人工记录。",
            ),
            (
                "预期目标",
                "1. 系统能够识别用户转人工意图。\n"
                "2. 前端能够提示已提交人工处理请求。\n"
                "3. 后台能够查看对应会话记录。\n"
                "4. 会话状态能够体现人工介入需求。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
    (
        "表5-8 知识库命中与会话管理功能测试用例",
        [
            ("测试用例编号", "TC08"),
            ("测试用例名称", "知识库命中与会话管理功能测试"),
            (
                "测试步骤",
                "1. 管理员在后台维护 AI 知识库内容。\n"
                "2. 用户提问与知识库内容相关的问题。\n"
                "3. 检查 AI 回复是否与知识库信息一致。\n"
                "4. 管理员进入后台查看会话记录与反馈情况。",
            ),
            (
                "预期目标",
                "1. 系统能够基于知识库内容生成回复。\n"
                "2. 回复结果与后台维护的信息保持一致。\n"
                "3. 会话记录能够在后台正常查看。\n"
                "4. 管理员能够对会话内容进行统一管理。",
            ),
            ("测试结果", "符合预期"),
        ],
    ),
]


RESULT_TABLE = [
    ("测试用例编号", "是否符合预期"),
    ("TC01", "是"),
    ("TC02", "是"),
    ("TC03", "是"),
    ("TC04", "是"),
    ("TC05", "是"),
    ("TC06", "是"),
    ("TC07", "是"),
    ("TC08", "是"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)


def format_paragraph(paragraph, bold=False, center=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        run.bold = bold


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_two_column_table(document, caption, rows):
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.font.name = "宋体"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption_run.font.size = Pt(10.5)
    caption_run.bold = True

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for left_text, right_text in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 4.2)
        set_cell_width(cells[1], 11.8)
        cells[0].text = left_text
        cells[1].text = right_text
        for idx, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, bold=(idx == 0), center=False)

    document.add_paragraph()


def add_result_table(document):
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("表5-9 测试结果表")
    caption_run.font.name = "宋体"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption_run.font.size = Pt(10.5)
    caption_run.bold = True

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for row_index, (left_text, right_text) in enumerate(RESULT_TABLE):
        cells = table.add_row().cells
        set_cell_width(cells[0], 8.0)
        set_cell_width(cells[1], 8.0)
        cells[0].text = left_text
        cells[1].text = right_text
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cells[0].paragraphs:
            format_paragraph(paragraph, bold=(row_index == 0), center=True)
        for paragraph in cells[1].paragraphs:
            format_paragraph(paragraph, bold=(row_index == 0), center=True)


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("第五章系统测试用例表")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(15)
    run.bold = True

    document.add_paragraph()

    for caption, rows in TEST_TABLES:
        add_two_column_table(document, caption, rows)

    add_result_table(document)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
