from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('表3-2 AI客服咨询用例规约')
r.bold = True
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.size = Pt(12)

rows = [
    ('用例编号', 'UC02'),
    ('用例名称', 'AI客服咨询'),
    ('简要说明', '该用例描述用户在小程序中通过AI客服发起咨询，系统根据问题内容进行意图识别，并结合知识库、业务数据查询及人工转接机制返回咨询结果。'),
    ('参与者', '用户'),
    ('前置条件', '用户已进入AI咨询页面；系统AI客服服务可正常使用。涉及订单、课包、优惠券等个人业务查询时，用户已登录。'),
    ('后置条件', '系统生成或更新会话记录，保存用户消息和回复消息；必要时生成转人工记录并更新会话状态。'),
    ('触发条件', '用户在AI咨询页面输入问题并发送消息后开始该用例。'),
    ('基本流', '1. 用户进入AI咨询页面。\n2. 用户输入咨询内容并发送消息。\n3. 系统校验消息内容，并根据当前用户标识或会话信息准备咨询会话。\n4. 系统保存用户发送的消息。\n5. 系统识别当前问题意图，并进入相应业务处理分支。\n6. 系统结合知识库内容、业务数据查询结果或规则模板生成回复内容。\n7. 系统保存回复消息，并更新会话摘要信息。\n8. 系统将咨询结果返回给用户显示。'),
    ('备选流', '3a. 用户输入内容为空：系统提示用户输入咨询内容，用例结束。\n5a. 用户咨询订单、课包、优惠券等个人业务信息但未登录：系统提示用户先登录后再查询，用例结束。\n6a. 用户请求人工客服或当前会话已进入人工处理状态：系统生成或保持转人工处理记录，系统不再进行AI自动回复，等待人工处理。'),
    ('成功场景', '用户成功提交咨询内容，系统返回对应回复，并完成会话记录更新；必要时系统成功转入人工处理。'),
    ('失败场景', '1. 用户输入内容为空，无法发起咨询。\n2. 用户未登录，无法查询个人业务信息。\n3. 当前会话转入人工处理后，系统不再继续自动回复。')
]

table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

table.columns[0].width = Cm(3.2)
table.columns[1].width = Cm(11.8)

for key, value in rows:
    cells = table.add_row().cells
    set_cell_text(cells[0], key, bold=True)
    set_cell_text(cells[1], value)
    set_cell_shading(cells[0], 'EDEDED')
    for c in cells:
        tc_pr = c._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in('w:tcW')
        if tc_w is None:
            tc_w = OxmlElement('w:tcW')
            tc_pr.append(tc_w)
        if c == cells[0]:
            tc_w.set(qn('w:w'), str(int(3.2 / 2.54 * 1440)))
        else:
            tc_w.set(qn('w:w'), str(int(11.8 / 2.54 * 1440)))
        tc_w.set(qn('w:type'), 'dxa')

# tighten paragraphs inside multi-line cells
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25

out = '/Users/atrox/集美大学/Graduation-Project/sports-course/docs/表3-2_AI客服咨询用例规约.docx'
doc.save(out)
print(out)
